"""Convert raw uploaded files into LangChain `Document` chunks.

Supports PDF (text + detected tables + optional OCR on embedded images),
DOCX (paragraphs + tables), TXT/MD, Excel (.xlsx/.xls), and raster images
(PNG/JPEG/WebP/GIF via OCR).
"""

from __future__ import annotations

import io
import logging
import os
import re
import subprocess
import tempfile
import shutil
from pathlib import Path
from typing import Any, List, Sequence

import fitz  # PyMuPDF
from langchain_core.documents import Document
from langchain_community.document_loaders import TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter

from backend.config import get_settings

logger = logging.getLogger(__name__)

_tesseract_missing_logged = False
_tesseract_exec_cached: str | None = None
_ocr_fail_once: str | None = None

# Detect figure captions (survey papers label diagrams “Figure 2: …” on the same page).
_FIGURE_CAPTION_RE = re.compile(r"\b(Figure|Fig\.)\s*\d+", re.IGNORECASE)

_SUPPORTED_EXT = {
    ".pdf",
    ".txt",
    ".md",
    ".docx",
    ".xlsx",
    ".xls",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".gif",
}


def _splitter() -> RecursiveCharacterTextSplitter:
    cfg = get_settings()
    return RecursiveCharacterTextSplitter(
        chunk_size=cfg.CHUNK_SIZE,
        chunk_overlap=cfg.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " ", ""],
        length_function=len,
    )


def _table_rows_to_markdown(rows: Sequence[Sequence[Any]]) -> str:
    """Turn tabular rows into a compact GitHub-flavored markdown table."""
    if not rows:
        return ""
    clean: List[List[str]] = []
    for row in rows:
        cells = [str(c).replace("|", "\\|").strip() if c is not None else "" for c in row]
        clean.append(cells)
    if not clean:
        return ""
    width = max(len(r) for r in clean)
    for r in clean:
        while len(r) < width:
            r.append("")
    header = clean[0]
    sep = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for body_row in clean[1:]:
        lines.append("| " + " | ".join(body_row) + " |")
    return "\n".join(lines)


def _resolve_tesseract_executable() -> str | None:
    """Find a tesseract binary that actually runs (`--version`), not just Path.exists.

    Startup previously logged “ready” when the file existed; pytesseract then failed
    with ENOENT/PATH errors when the subprocess could not exec the binary (sandbox,
    permissions, or strip-needed env). subprocess.run is the source of truth.
    """
    global _tesseract_exec_cached
    if _tesseract_exec_cached:
        return _tesseract_exec_cached

    cfg = get_settings()
    candidates: List[str] = []
    if cfg.TESSERACT_CMD:
        candidates.append(cfg.TESSERACT_CMD)
    w = shutil.which("tesseract")
    if w:
        candidates.append(w)
    candidates.extend(
        (
            "/opt/homebrew/bin/tesseract",
            "/usr/local/bin/tesseract",
            "/usr/local/opt/tesseract/bin/tesseract",
        )
    )

    seen: set[str] = set()
    for raw in candidates:
        if not raw or raw in seen:
            continue
        seen.add(raw)
        path = Path(raw).expanduser()
        try:
            path = path.resolve()
        except OSError:
            continue
        if not path.is_file():
            continue
        if not os.access(path, os.X_OK):
            logger.debug("Skipping non-executable Tesseract candidate: %s", path)
            continue
        try:
            proc = subprocess.run(
                [str(path), "--version"],
                capture_output=True,
                text=True,
                timeout=25,
                env=os.environ.copy(),
            )
        except OSError as exc:
            logger.debug("Cannot exec %s: %s", path, exc)
            continue
        if proc.returncode != 0:
            logger.debug(
                "tesseract --version failed at %s rc=%s stderr=%s",
                path,
                proc.returncode,
                (proc.stderr or "")[:500],
            )
            continue

        _tesseract_exec_cached = str(path)
        logger.debug("Verified Tesseract at %s", _tesseract_exec_cached)
        return _tesseract_exec_cached

    return None


def _ensure_tesseract() -> bool:
    """Configure pytesseract with a verified executable path."""
    global _tesseract_missing_logged
    if not get_settings().DOCUMENT_OCR_ENABLED:
        return False
    try:
        import pytesseract
    except ImportError:
        if not _tesseract_missing_logged:
            logger.warning("Install pytesseract + Pillow for OCR: pip install pytesseract Pillow")
            _tesseract_missing_logged = True
        return False

    resolved = _resolve_tesseract_executable()
    if not resolved:
        if not _tesseract_missing_logged:
            logger.warning(
                "No working Tesseract binary found (ran --version on PATH and "
                "common Homebrew paths). Install: brew install tesseract then set "
                "TESSERACT_CMD to the output of: which tesseract"
            )
            _tesseract_missing_logged = True
        return False

    pytesseract.pytesseract.tesseract_cmd = resolved
    return True


def _ocr_pil(img: Any, *, tess_config: str) -> str:
    global _ocr_fail_once
    import pytesseract
    from PIL import Image

    if not isinstance(img, Image.Image):
        img = img.convert("RGB")
    w, h = img.size
    max_side = 2600
    if max(w, h) > max_side:
        scale = max_side / max(w, h)
        img = img.resize(
            (max(1, int(w * scale)), max(1, int(h * scale))),
            Image.Resampling.LANCZOS,
        )
    try:
        return (pytesseract.image_to_string(img, config=tess_config) or "").strip()
    except Exception as exc:  # noqa: BLE001
        msg = str(exc)
        if _ocr_fail_once != msg:
            logger.warning(
                "Tesseract OCR failed: %s "
                "(pytesseract binary=%s — run `tesseract --version` in Terminal; "
                "set TESSERACT_CMD in .env to that exact path)",
                exc,
                pytesseract.pytesseract.tesseract_cmd,
            )
            _ocr_fail_once = msg
        else:
            logger.debug("Tesseract OCR repeat failure suppressed: %s", exc)
        return ""


def _ocr_bytes(image_bytes: bytes, *, tess_config: str | None = None) -> str:
    if not get_settings().DOCUMENT_OCR_ENABLED:
        return ""
    if not _ensure_tesseract():
        return ""
    cfg = get_settings()
    tc = tess_config or cfg.OCR_TESSERACT_CONFIG_IMAGE
    try:
        from PIL import Image

        img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        return _ocr_pil(img, tess_config=tc)
    except Exception as exc:  # noqa: BLE001
        logger.warning("OCR decode failed: %s", exc)
        return ""


def _merge_ocr_passes(primary: str, secondary: str) -> str:
    """Combine two OCR layouts; drop duplicate lines (case-insensitive)."""
    lines_out: List[str] = []
    seen: set[str] = set()
    for block in (primary, secondary):
        for line in block.splitlines():
            s = line.strip()
            if not s:
                continue
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            lines_out.append(s)
    return "\n".join(lines_out)


def _ocr_page_pixmap(page: fitz.Page, dpi: int) -> str:
    """Rasterize page and OCR — dual layout passes for diagram callouts + small fonts."""
    if not _ensure_tesseract():
        return ""
    cfg = get_settings()
    try:
        from PIL import Image

        scale = dpi / 72.0
        mat = fitz.Matrix(scale, scale)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        sparse = _ocr_pil(img, tess_config=cfg.OCR_TESSERACT_CONFIG_PAGE)
        if not cfg.PDF_PAGE_OCR_SECOND_PASS:
            return sparse
        blocky = _ocr_pil(img, tess_config=cfg.OCR_TESSERACT_CONFIG_PAGE_BLOCK)
        return _merge_ocr_passes(sparse, blocky)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Page pixmap OCR failed: %s", exc)
        return ""


def _load_pdf(path: Path) -> List[Document]:
    """Extract page text, tables, OCR from embedded bitmaps, and pixmap OCR for vector figures."""
    cfg = get_settings()
    doc = fitz.open(path)
    out: List[Document] = []
    try:
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            parts: List[str] = []

            body_text = (page.get_text("text") or "").strip()
            if body_text:
                parts.append(body_text)

            mentions_figure = bool(_FIGURE_CAPTION_RE.search(body_text))

            try:
                tf = page.find_tables()
                for tab in getattr(tf, "tables", []) or []:
                    extracted = tab.extract()
                    if extracted:
                        md = _table_rows_to_markdown(extracted)
                        if md:
                            parts.append(
                                f"[Table on PDF page {page_index + 1}]\n{md}"
                            )
            except Exception as exc:  # noqa: BLE001
                logger.debug("Table extraction skipped page %s: %s", page_index, exc)

            embedded_chars = 0
            page_imgs: List[Any] = []
            try:
                page_imgs = page.get_images(full=True) or []
                for img in page_imgs:
                    xref = img[0]
                    info = doc.extract_image(xref)
                    img_bytes = info.get("image") or b""
                    if not img_bytes:
                        continue
                    ocr_text = _ocr_bytes(img_bytes)
                    if ocr_text:
                        embedded_chars += len(ocr_text)
                        parts.append(
                            f"[Text from embedded image on PDF page {page_index + 1}]\n{ocr_text}"
                        )
            except Exception as exc:  # noqa: BLE001
                logger.debug("Embedded image OCR skipped page %s: %s", page_index, exc)

            # Vector diagrams: rasterize page. Force OCR when caption references Figure N (survey PDFs).
            if cfg.PDF_PAGE_RENDER_OCR and cfg.DOCUMENT_OCR_ENABLED:
                no_bitmaps = len(page_imgs) == 0
                weak_embedded = embedded_chars < cfg.PDF_MIN_EMBEDDED_OCR_CHARS
                force_figure_page = (
                    cfg.PDF_PAGE_OCR_ON_FIGURE_CAPTION and mentions_figure
                )
                if no_bitmaps or weak_embedded or force_figure_page:
                    visual = _ocr_page_pixmap(page, cfg.PDF_PAGE_OCR_DPI)
                    if visual.strip():
                        parts.append(
                            f"[Page {page_index + 1} — text from figures/diagrams (OCR)]\n{visual}"
                        )

            merged = "\n\n".join(p for p in parts if p)
            if not merged.strip():
                merged = f"[Page {page_index + 1} — no extractable text]"
            out.append(
                Document(
                    page_content=merged,
                    metadata={"page": page_index},
                )
            )
    finally:
        doc.close()
    return out


def _load_docx(path: Path) -> List[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts: List[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for ti, table in enumerate(doc.tables):
        rows: List[List[str]] = []
        for row in table.rows:
            rows.append([(cell.text or "").strip() for cell in row.cells])
        if rows:
            md = _table_rows_to_markdown(rows)
            if md:
                parts.append(f"[Word table {ti + 1}]\n{md}")
    body = "\n\n".join(parts)
    if not body.strip():
        body = "[No text or tables found in document]"
    return [Document(page_content=body, metadata={"page": 0})]


def _load_excel(path: Path) -> List[Document]:
    ext = path.suffix.lower()
    docs: List[Document] = []

    if ext == ".xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(path, read_only=True, data_only=True)
        try:
            for sheet in wb.worksheets:
                rows: List[List[str]] = []
                for row in sheet.iter_rows(values_only=True):
                    rows.append(
                        [str(c) if c is not None else "" for c in row]
                    )
                rows = [r for r in rows if any(x.strip() for x in r)]
                if not rows:
                    continue
                md = _table_rows_to_markdown(rows)
                title = sheet.title or "Sheet"
                block = f"## Excel sheet: {title}\n\n{md}"
                docs.append(Document(page_content=block, metadata={"page": 0}))
        finally:
            wb.close()

    elif ext == ".xls":
        import xlrd

        book = xlrd.open_workbook(str(path))
        for sheet in book.sheets():
            rows = []
            for r in range(sheet.nrows):
                rows.append([str(sheet.cell_value(r, c)) for c in range(sheet.ncols)])
            rows = [r for r in rows if any(x.strip() for x in r)]
            if not rows:
                continue
            md = _table_rows_to_markdown(rows)
            title = sheet.name or "Sheet"
            block = f"## Excel sheet: {title}\n\n{md}"
            docs.append(Document(page_content=block, metadata={"page": 0}))

    if not docs:
        raise ValueError("Excel file contained no non-empty sheets")
    return docs


def _load_plain(path: Path) -> List[Document]:
    return TextLoader(str(path), encoding="utf-8").load()


def _load_image(path: Path) -> List[Document]:
    raw = path.read_bytes()
    text = _ocr_bytes(raw)
    if not text.strip():
        raise ValueError(
            "No text extracted from image. Install Tesseract (e.g. `brew install "
            "tesseract`) and ensure DOCUMENT_OCR_ENABLED=true, or try a clearer scan."
        )
    meta = (
        "This chunk came from an image file; text was produced via OCR.\n\n"
        f"{text}"
    )
    return [Document(page_content=meta, metadata={"page": 0})]


def _load(path: Path) -> List[Document]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _load_pdf(path)
    if ext == ".docx":
        return _load_docx(path)
    if ext in {".xlsx", ".xls"}:
        return _load_excel(path)
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
        return _load_image(path)
    if ext in {".txt", ".md"}:
        return _load_plain(path)
    raise ValueError(f"Unsupported file extension: {ext}")


def log_ocr_status_at_startup() -> None:
    """Called once from FastAPI lifespan — surfaces PATH/Tesseract issues early."""
    cfg = get_settings()
    if not cfg.DOCUMENT_OCR_ENABLED:
        logger.info("DOCUMENT_OCR_ENABLED=false — OCR skipped for PDFs/images")
        return
    if _ensure_tesseract():
        import pytesseract

        cmd = getattr(pytesseract.pytesseract, "tesseract_cmd", "tesseract")
        logger.info(
            "Tesseract OCR ready (%s); page-render fallback=%s dpi=%d",
            cmd,
            cfg.PDF_PAGE_RENDER_OCR,
            cfg.PDF_PAGE_OCR_DPI,
        )
    else:
        logger.warning(
            "Tesseract not available — figure/diagram text in PDFs will be missing until configured."
        )


def parse(file_bytes: bytes, filename: str) -> List[Document]:
    """Parse raw bytes into chunked `Document` objects with citation metadata."""
    ext = Path(filename).suffix.lower()
    if ext not in _SUPPORTED_EXT:
        raise ValueError(
            f"Unsupported file type '{ext}'. Allowed: {sorted(_SUPPORTED_EXT)}"
        )

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(file_bytes)
        tmp_path = Path(tmp.name)

    try:
        raw_docs = _load(tmp_path)
        chunks = _splitter().split_documents(raw_docs)
    finally:
        tmp_path.unlink(missing_ok=True)

    enriched: List[Document] = []
    for idx, chunk in enumerate(chunks):
        page_raw = chunk.metadata.get("page")
        # Loaders use 0-based PDF indices or 0 for single-part docs → store 1-based for citations.
        if isinstance(page_raw, (int, float)):
            page_display = int(page_raw) + 1
        else:
            page_display = 0
        enriched.append(
            Document(
                page_content=chunk.page_content,
                metadata={
                    "source": filename,
                    "page": page_display,
                    "chunk_index": idx,
                },
            )
        )
    return enriched
